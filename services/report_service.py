import os
import io
import pandas as pd
from datetime import datetime
from decimal import Decimal

# Para relatórios Excel
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

# Para relatórios Word
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_invoice_report(invoices, format_type, output_dir, include_items=False):
    """
    Gera um relatório de notas fiscais
    
    Args:
        invoices: Lista de objetos Invoice
        format_type: Tipo de formato ('excel' ou 'word')
        output_dir: Diretório para salvar o arquivo
        include_items: Se True, inclui os itens das notas
    
    Returns:
        Caminho do arquivo gerado
    """
    if format_type == 'excel':
        return _generate_invoice_excel(invoices, output_dir, include_items)
    else:
        return _generate_invoice_word(invoices, output_dir, include_items)


def generate_user_report(users, format_type, output_dir):
    """
    Gera um relatório de usuários
    
    Args:
        users: Lista de objetos User
        format_type: Tipo de formato ('excel' ou 'word')
        output_dir: Diretório para salvar o arquivo
    
    Returns:
        Caminho do arquivo gerado
    """
    if format_type == 'excel':
        return _generate_user_excel(users, output_dir)
    else:
        return _generate_user_word(users, output_dir)


def generate_inventory_report(products, format_type, output_dir):
    """
    Gera um relatório de estoque
    
    Args:
        products: Lista de objetos Product
        format_type: Tipo de formato ('excel' ou 'word')
        output_dir: Diretório para salvar o arquivo
    
    Returns:
        Caminho do arquivo gerado
    """
    if format_type == 'excel':
        return _generate_inventory_excel(products, output_dir)
    else:
        return _generate_inventory_word(products, output_dir)


def generate_financial_report(invoices, format_type, output_dir, start_date, end_date):
    """
    Gera um relatório financeiro
    
    Args:
        invoices: Lista de objetos Invoice
        format_type: Tipo de formato ('excel' ou 'word')
        output_dir: Diretório para salvar o arquivo
        start_date: Data inicial do período
        end_date: Data final do período
    
    Returns:
        Caminho do arquivo gerado
    """
    if format_type == 'excel':
        return _generate_financial_excel(invoices, output_dir, start_date, end_date)
    else:
        return _generate_financial_word(invoices, output_dir, start_date, end_date)


# Funções privadas para geração de relatórios em Excel
def _generate_invoice_excel(invoices, output_dir, include_items):
    """Gera relatório de notas fiscais em Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Notas Fiscais"
    
    # Definir cabeçalhos
    headers = [
        'Número', 'Série', 'Tipo', 'Status', 'Data Emissão', 
        'Cliente/Fornecedor', 'Valor Total', 'Impostos', 'Criado em'
    ]
    
    # Estilo para cabeçalhos
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    # Adicionar cabeçalhos
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    # Adicionar dados
    for row, invoice in enumerate(invoices, 2):
        ws.cell(row=row, column=1, value=invoice.invoice_number)
        ws.cell(row=row, column=2, value=invoice.series)
        ws.cell(row=row, column=3, value=invoice.type.value)
        ws.cell(row=row, column=4, value=invoice.status.value)
        ws.cell(row=row, column=5, value=invoice.issue_date.strftime('%d/%m/%Y'))
        
        # Cliente ou fornecedor
        entity_name = ""
        if invoice.customer_id and invoice.customer:
            entity_name = invoice.customer.name
        elif invoice.supplier_id and invoice.supplier:
            entity_name = invoice.supplier.name
        ws.cell(row=row, column=6, value=entity_name)
        
        ws.cell(row=row, column=7, value=float(invoice.total_value))
        ws.cell(row=row, column=8, value=float(invoice.total_tax))
        ws.cell(row=row, column=9, value=invoice.created_at.strftime('%d/%m/%Y %H:%M'))
    
    # Formatar colunas
    ws.column_dimensions['A'].width = 15  # Número
    ws.column_dimensions['B'].width = 8   # Série
    ws.column_dimensions['C'].width = 12  # Tipo
    ws.column_dimensions['D'].width = 12  # Status
    ws.column_dimensions['E'].width = 15  # Data Emissão
    ws.column_dimensions['F'].width = 30  # Cliente/Fornecedor
    ws.column_dimensions['G'].width = 15  # Valor Total
    ws.column_dimensions['H'].width = 15  # Impostos
    ws.column_dimensions['I'].width = 20  # Criado em
    
    # Formatar valores monetários
    for row in range(2, len(invoices) + 2):
        ws.cell(row=row, column=7).number_format = 'R$ #,##0.00'
        ws.cell(row=row, column=8).number_format = 'R$ #,##0.00'
    
    # Se deve incluir itens, adicionar uma segunda planilha
    if include_items:
        ws_items = wb.create_sheet("Itens das Notas")
        
        # Cabeçalhos para itens
        item_headers = [
            'Nota Fiscal', 'Série', 'Produto', 'Quantidade', 
            'Preço Unitário', 'Desconto', 'Total'
        ]
        
        for col, header in enumerate(item_headers, 1):
            cell = ws_items.cell(row=1, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
        
        # Adicionar itens
        row = 2
        for invoice in invoices:
            for item in invoice.items:
                ws_items.cell(row=row, column=1, value=invoice.invoice_number)
                ws_items.cell(row=row, column=2, value=invoice.series)
                ws_items.cell(row=row, column=3, value=item.product.name if item.product else "")
                ws_items.cell(row=row, column=4, value=float(item.quantity))
                ws_items.cell(row=row, column=5, value=float(item.unit_price))
                ws_items.cell(row=row, column=6, value=float(item.discount))
                ws_items.cell(row=row, column=7, value=float(item.total))
                row += 1
        
        # Formatar colunas de itens
        ws_items.column_dimensions['A'].width = 15  # Nota Fiscal
        ws_items.column_dimensions['B'].width = 8   # Série
        ws_items.column_dimensions['C'].width = 30  # Produto
        ws_items.column_dimensions['D'].width = 12  # Quantidade
        ws_items.column_dimensions['E'].width = 15  # Preço Unitário
        ws_items.column_dimensions['F'].width = 12  # Desconto
        ws_items.column_dimensions['G'].width = 15  # Total
        
        # Formatar valores monetários
        for r in range(2, row):
            ws_items.cell(row=r, column=5).number_format = 'R$ #,##0.00'
            ws_items.cell(row=r, column=6).number_format = 'R$ #,##0.00'
            ws_items.cell(row=r, column=7).number_format = 'R$ #,##0.00'
    
    # Salvar arquivo
    filename = os.path.join(output_dir, f"Relatorio_Notas_Fiscais_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    wb.save(filename)
    
    return filename


def _generate_user_excel(users, output_dir):
    """Gera relatório de usuários em Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Usuários"
    
    # Definir cabeçalhos
    headers = [
        'ID', 'Nome de Usuário', 'Nome Completo', 'Email', 
        'Função', 'Ativo', 'Criado em', 'Último Acesso'
    ]
    
    # Estilo para cabeçalhos
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    # Adicionar cabeçalhos
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    # Adicionar dados
    for row, user in enumerate(users, 2):
        ws.cell(row=row, column=1, value=user.id)
        ws.cell(row=row, column=2, value=user.username)
        ws.cell(row=row, column=3, value=user.full_name)
        ws.cell(row=row, column=4, value=user.email)
        ws.cell(row=row, column=5, value=user.role.value if user.role else "")
        ws.cell(row=row, column=6, value="Sim" if user.is_active else "Não")
        ws.cell(row=row, column=7, value=user.created_at.strftime('%d/%m/%Y %H:%M') if user.created_at else "")
        ws.cell(row=row, column=8, value=user.last_login.strftime('%d/%m/%Y %H:%M') if user.last_login else "Nunca")
    
    # Formatar colunas
    ws.column_dimensions['A'].width = 6   # ID
    ws.column_dimensions['B'].width = 20  # Nome de Usuário
    ws.column_dimensions['C'].width = 30  # Nome Completo
    ws.column_dimensions['D'].width = 30  # Email
    ws.column_dimensions['E'].width = 15  # Função
    ws.column_dimensions['F'].width = 10  # Ativo
    ws.column_dimensions['G'].width = 20  # Criado em
    ws.column_dimensions['H'].width = 20  # Último Acesso
    
    # Salvar arquivo
    filename = os.path.join(output_dir, f"Relatorio_Usuarios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    wb.save(filename)
    
    return filename


def _generate_inventory_excel(products, output_dir):
    """Gera relatório de estoque em Excel"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Estoque"
    
    # Definir cabeçalhos
    headers = [
        'SKU', 'Nome do Produto', 'Descrição', 
        'Preço de Compra', 'Preço de Venda', 
        'Estoque Atual', 'Estoque Mínimo', 
        'NCM', 'Peso (kg)'
    ]
    
    # Estilo para cabeçalhos
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    # Adicionar cabeçalhos
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    # Adicionar dados
    for row, product in enumerate(products, 2):
        current_stock = product.current_stock()
        
        ws.cell(row=row, column=1, value=product.sku)
        ws.cell(row=row, column=2, value=product.name)
        ws.cell(row=row, column=3, value=product.description)
        ws.cell(row=row, column=4, value=float(product.purchase_price))
        ws.cell(row=row, column=5, value=float(product.sale_price))
        ws.cell(row=row, column=6, value=current_stock)
        ws.cell(row=row, column=7, value=product.min_stock)
        ws.cell(row=row, column=8, value=product.ncm)
        ws.cell(row=row, column=9, value=float(product.weight) if product.weight else None)
        
        # Destacar estoque baixo
        if current_stock is not None and current_stock <= product.min_stock:
            for col in range(1, 10):
                ws.cell(row=row, column=col).fill = PatternFill(start_color="FFAAAA", end_color="FFAAAA", fill_type="solid")
    
    # Formatar colunas
    ws.column_dimensions['A'].width = 15  # SKU
    ws.column_dimensions['B'].width = 30  # Nome do Produto
    ws.column_dimensions['C'].width = 40  # Descrição
    ws.column_dimensions['D'].width = 15  # Preço de Compra
    ws.column_dimensions['E'].width = 15  # Preço de Venda
    ws.column_dimensions['F'].width = 15  # Estoque Atual
    ws.column_dimensions['G'].width = 15  # Estoque Mínimo
    ws.column_dimensions['H'].width = 10  # NCM
    ws.column_dimensions['I'].width = 10  # Peso
    
    # Formatar valores monetários
    for row in range(2, len(products) + 2):
        ws.cell(row=row, column=4).number_format = 'R$ #,##0.00'
        ws.cell(row=row, column=5).number_format = 'R$ #,##0.00'
    
    # Salvar arquivo
    filename = os.path.join(output_dir, f"Relatorio_Estoque_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    wb.save(filename)
    
    return filename


def _generate_financial_excel(invoices, output_dir, start_date, end_date):
    """Gera relatório financeiro em Excel"""
    wb = Workbook()
    
    # Planilha de resumo
    ws_summary = wb.active
    ws_summary.title = "Resumo Financeiro"
    
    # Adicionar período do relatório
    ws_summary.cell(row=1, column=1, value="Relatório Financeiro")
    ws_summary.cell(row=1, column=1).font = Font(size=16, bold=True)
    
    ws_summary.cell(row=2, column=1, value=f"Período: {start_date.strftime('%d/%m/%Y')} a {end_date.strftime('%d/%m/%Y')}")
    ws_summary.cell(row=2, column=1).font = Font(size=12)
    
    # Separar notas de entrada e saída
    inbound_invoices = [inv for inv in invoices if inv.type.value == 'Entrada']
    outbound_invoices = [inv for inv in invoices if inv.type.value == 'Saída']
    
    # Calcular totais
    total_inbound = sum(float(inv.total_value) for inv in inbound_invoices)
    total_outbound = sum(float(inv.total_value) for inv in outbound_invoices)
    balance = total_outbound - total_inbound
    
    # Adicionar resumo
    row = 4
    ws_summary.cell(row=row, column=1, value="Resumo:")
    ws_summary.cell(row=row, column=1).font = Font(bold=True)
    
    row += 1
    ws_summary.cell(row=row, column=1, value="Total de Entradas:")
    ws_summary.cell(row=row, column=2, value=total_inbound)
    ws_summary.cell(row=row, column=2).number_format = 'R$ #,##0.00'
    
    row += 1
    ws_summary.cell(row=row, column=1, value="Total de Saídas:")
    ws_summary.cell(row=row, column=2, value=total_outbound)
    ws_summary.cell(row=row, column=2).number_format = 'R$ #,##0.00'
    
    row += 1
    ws_summary.cell(row=row, column=1, value="Saldo:")
    ws_summary.cell(row=row, column=2, value=balance)
    ws_summary.cell(row=row, column=2).number_format = 'R$ #,##0.00'
    
    # Estilo para valores positivos/negativos
    if balance >= 0:
        ws_summary.cell(row=row, column=2).font = Font(color="00AA00", bold=True)
    else:
        ws_summary.cell(row=row, column=2).font = Font(color="AA0000", bold=True)
    
    # Planilha detalhada de notas
    ws_invoices = wb.create_sheet("Notas Fiscais")
    
    # Definir cabeçalhos
    headers = [
        'Número', 'Série', 'Tipo', 'Data Emissão', 
        'Cliente/Fornecedor', 'Valor Total', 'Impostos'
    ]
    
    # Estilo para cabeçalhos
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    # Adicionar cabeçalhos
    for col, header in enumerate(headers, 1):
        cell = ws_invoices.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
    
    # Adicionar dados
    for row, invoice in enumerate(sorted(invoices, key=lambda x: x.issue_date), 2):
        ws_invoices.cell(row=row, column=1, value=invoice.invoice_number)
        ws_invoices.cell(row=row, column=2, value=invoice.series)
        ws_invoices.cell(row=row, column=3, value=invoice.type.value)
        ws_invoices.cell(row=row, column=4, value=invoice.issue_date.strftime('%d/%m/%Y'))
        
        # Cliente ou fornecedor
        entity_name = ""
        if invoice.customer_id and invoice.customer:
            entity_name = invoice.customer.name
        elif invoice.supplier_id and invoice.supplier:
            entity_name = invoice.supplier.name
        ws_invoices.cell(row=row, column=5, value=entity_name)
        
        ws_invoices.cell(row=row, column=6, value=float(invoice.total_value))
        ws_invoices.cell(row=row, column=7, value=float(invoice.total_tax))
        
        # Estilo diferente para entradas e saídas
        if invoice.type.value == 'Entrada':
            fill = PatternFill(start_color="FFDDDD", end_color="FFDDDD", fill_type="solid")
            for col in range(1, 8):
                ws_invoices.cell(row=row, column=col).fill = fill
    
    # Formatar colunas
    ws_invoices.column_dimensions['A'].width = 15  # Número
    ws_invoices.column_dimensions['B'].width = 8   # Série
    ws_invoices.column_dimensions['C'].width = 12  # Tipo
    ws_invoices.column_dimensions['D'].width = 15  # Data Emissão
    ws_invoices.column_dimensions['E'].width = 30  # Cliente/Fornecedor
    ws_invoices.column_dimensions['F'].width = 15  # Valor Total
    ws_invoices.column_dimensions['G'].width = 15  # Impostos
    
    # Formatar valores monetários
    for row in range(2, len(invoices) + 2):
        ws_invoices.cell(row=row, column=6).number_format = 'R$ #,##0.00'
        ws_invoices.cell(row=row, column=7).number_format = 'R$ #,##0.00'
    
    # Adicionar planilha de impostos
    from services.tax_service import get_tax_summary_by_period
    tax_summary = get_tax_summary_by_period(start_date, end_date)
    
    if tax_summary:
        ws_taxes = wb.create_sheet("Resumo de Impostos")
        
        # Cabeçalhos para impostos
        ws_taxes.cell(row=1, column=1, value="Imposto")
        ws_taxes.cell(row=1, column=2, value="Valor Total")
        
        # Estilo para cabeçalhos
        ws_taxes.cell(row=1, column=1).fill = header_fill
        ws_taxes.cell(row=1, column=1).font = header_font
        ws_taxes.cell(row=1, column=2).fill = header_fill
        ws_taxes.cell(row=1, column=2).font = header_font
        
        # Adicionar dados de impostos
        row = 2
        for tax_type, data in tax_summary.items():
            ws_taxes.cell(row=row, column=1, value=tax_type)
            ws_taxes.cell(row=row, column=2, value=data['total'])
            ws_taxes.cell(row=row, column=2).number_format = 'R$ #,##0.00'
            
            # Destacar valores negativos (créditos)
            if data['total'] < 0:
                ws_taxes.cell(row=row, column=2).font = Font(color="AA0000")
            
            row += 1
        
        # Formatar colunas
        ws_taxes.column_dimensions['A'].width = 20  # Imposto
        ws_taxes.column_dimensions['B'].width = 20  # Valor Total
    
    # Salvar arquivo
    filename = os.path.join(output_dir, f"Relatorio_Financeiro_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
    wb.save(filename)
    
    return filename


# Funções privadas para geração de relatórios em Word
def _generate_invoice_word(invoices, output_dir, include_items):
    """Gera relatório de notas fiscais em Word"""
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo de título
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Estilo de cabeçalho
    heading_style = styles['Heading 1']
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True
    
    # Título do documento
    doc.add_paragraph("Relatório de Notas Fiscais", style='Title')
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()
    
    # Tabela de notas fiscais
    doc.add_heading("Notas Fiscais", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    
    # Cabeçalhos
    header_cells = table.rows[0].cells
    header_cells[0].text = "Número"
    header_cells[1].text = "Série"
    header_cells[2].text = "Tipo"
    header_cells[3].text = "Status"
    header_cells[4].text = "Data Emissão"
    header_cells[5].text = "Cliente/Fornecedor"
    header_cells[6].text = "Valor Total"
    
    # Formatar cabeçalhos
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
    
    # Dados
    for invoice in invoices:
        row_cells = table.add_row().cells
        row_cells[0].text = invoice.invoice_number
        row_cells[1].text = invoice.series
        row_cells[2].text = invoice.type.value
        row_cells[3].text = invoice.status.value
        row_cells[4].text = invoice.issue_date.strftime('%d/%m/%Y')
        
        # Cliente ou fornecedor
        entity_name = ""
        if invoice.customer_id and hasattr(invoice, 'customer') and invoice.customer:
            entity_name = invoice.customer.name
        elif invoice.supplier_id and hasattr(invoice, 'supplier') and invoice.supplier:
            entity_name = invoice.supplier.name
        row_cells[5].text = entity_name
        
        row_cells[6].text = f"R$ {float(invoice.total_value):,.2f}"
    
    # Se deve incluir itens, adicionar seção de itens
    if include_items:
        doc.add_page_break()
        doc.add_heading("Itens das Notas Fiscais", level=1)
        
        for invoice in invoices:
            doc.add_heading(f"Nota: {invoice.invoice_number} - {invoice.series}", level=2)
            
            # Tabela de itens
            if invoice.items.count() > 0:
                table_items = doc.add_table(rows=1, cols=5)
                table_items.style = 'Table Grid'
                
                # Cabeçalhos
                header_cells = table_items.rows[0].cells
                header_cells[0].text = "Produto"
                header_cells[1].text = "Quantidade"
                header_cells[2].text = "Preço Unitário"
                header_cells[3].text = "Desconto"
                header_cells[4].text = "Total"
                
                # Formatar cabeçalhos
                for cell in header_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell.paragraphs[0].runs[0]
                    run.font.bold = True
                
                # Dados dos itens
                for item in invoice.items:
                    row_cells = table_items.add_row().cells
                    row_cells[0].text = item.product.name if item.product else ""
                    row_cells[1].text = str(float(item.quantity))
                    row_cells[2].text = f"R$ {float(item.unit_price):,.2f}"
                    row_cells[3].text = f"R$ {float(item.discount):,.2f}"
                    row_cells[4].text = f"R$ {float(item.total):,.2f}"
            else:
                doc.add_paragraph("Sem itens para esta nota.")
            
            doc.add_paragraph()
    
    # Salvar o documento
    filename = os.path.join(output_dir, f"Relatorio_Notas_Fiscais_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    
    return filename


def _generate_user_word(users, output_dir):
    """Gera relatório de usuários em Word"""
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo de título
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Título do documento
    doc.add_paragraph("Relatório de Usuários", style='Title')
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()
    
    # Tabela de usuários
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Cabeçalhos
    header_cells = table.rows[0].cells
    header_cells[0].text = "Nome de Usuário"
    header_cells[1].text = "Nome Completo"
    header_cells[2].text = "Email"
    header_cells[3].text = "Função"
    header_cells[4].text = "Ativo"
    header_cells[5].text = "Último Acesso"
    
    # Formatar cabeçalhos
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
    
    # Dados
    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = user.username
        row_cells[1].text = user.full_name
        row_cells[2].text = user.email
        row_cells[3].text = user.role.value if user.role else ""
        row_cells[4].text = "Sim" if user.is_active else "Não"
        row_cells[5].text = user.last_login.strftime('%d/%m/%Y %H:%M') if user.last_login else "Nunca"
    
    # Salvar o documento
    filename = os.path.join(output_dir, f"Relatorio_Usuarios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    
    return filename


def _generate_inventory_word(products, output_dir):
    """Gera relatório de estoque em Word"""
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo de título
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Título do documento
    doc.add_paragraph("Relatório de Estoque", style='Title')
    doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()
    
    # Tabela de produtos
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Cabeçalhos
    header_cells = table.rows[0].cells
    header_cells[0].text = "SKU"
    header_cells[1].text = "Nome do Produto"
    header_cells[2].text = "Preço de Compra"
    header_cells[3].text = "Preço de Venda"
    header_cells[4].text = "Estoque Atual"
    header_cells[5].text = "Estoque Mínimo"
    
    # Formatar cabeçalhos
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
    
    # Dados
    for product in products:
        current_stock = product.current_stock()
        
        row_cells = table.add_row().cells
        row_cells[0].text = product.sku
        row_cells[1].text = product.name
        row_cells[2].text = f"R$ {float(product.purchase_price):,.2f}"
        row_cells[3].text = f"R$ {float(product.sale_price):,.2f}"
        row_cells[4].text = str(current_stock)
        row_cells[5].text = str(product.min_stock)
        
        # Destacar produtos com estoque baixo
        if current_stock is not None and current_stock <= product.min_stock:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Salvar o documento
    filename = os.path.join(output_dir, f"Relatorio_Estoque_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    
    return filename


def _generate_financial_word(invoices, output_dir, start_date, end_date):
    """Gera relatório financeiro em Word"""
    doc = Document()
    
    # Configurar estilos
    styles = doc.styles
    
    # Estilo de título
    title_style = styles['Title']
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    
    # Título do documento
    doc.add_paragraph("Relatório Financeiro", style='Title')
    doc.add_paragraph(f"Período: {start_date.strftime('%d/%m/%Y')} a {end_date.strftime('%d/%m/%Y')}")
    doc.add_paragraph()
    
    # Separar notas de entrada e saída
    inbound_invoices = [inv for inv in invoices if inv.type.value == 'Entrada']
    outbound_invoices = [inv for inv in invoices if inv.type.value == 'Saída']
    
    # Calcular totais
    total_inbound = sum(float(inv.total_value) for inv in inbound_invoices)
    total_outbound = sum(float(inv.total_value) for inv in outbound_invoices)
    balance = total_outbound - total_inbound
    
    # Resumo financeiro
    doc.add_heading("Resumo Financeiro", level=1)
    
    p = doc.add_paragraph("Total de Entradas: ")
    p.add_run(f"R$ {total_inbound:,.2f}").bold = True
    
    p = doc.add_paragraph("Total de Saídas: ")
    p.add_run(f"R$ {total_outbound:,.2f}").bold = True
    
    p = doc.add_paragraph("Saldo: ")
    run = p.add_run(f"R$ {balance:,.2f}")
    run.bold = True
    if balance >= 0:
        run.font.color.rgb = RGBColor(0, 128, 0)  # Verde para saldo positivo
    else:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Vermelho para saldo negativo
    
    doc.add_paragraph()
    
    # Detalhamento de notas fiscais
    doc.add_heading("Detalhamento de Notas Fiscais", level=1)
    
    # Notas de saída
    if outbound_invoices:
        doc.add_heading("Notas de Saída", level=2)
        
        table_out = doc.add_table(rows=1, cols=5)
        table_out.style = 'Table Grid'
        
        # Cabeçalhos
        header_cells = table_out.rows[0].cells
        header_cells[0].text = "Número"
        header_cells[1].text = "Data Emissão"
        header_cells[2].text = "Cliente"
        header_cells[3].text = "Valor Total"
        header_cells[4].text = "Impostos"
        
        # Formatar cabeçalhos
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
        
        # Dados
        for invoice in sorted(outbound_invoices, key=lambda x: x.issue_date):
            row_cells = table_out.add_row().cells
            row_cells[0].text = f"{invoice.invoice_number}-{invoice.series}"
            row_cells[1].text = invoice.issue_date.strftime('%d/%m/%Y')
            row_cells[2].text = invoice.customer.name if invoice.customer else ""
            row_cells[3].text = f"R$ {float(invoice.total_value):,.2f}"
            row_cells[4].text = f"R$ {float(invoice.total_tax):,.2f}"
    
    # Notas de entrada
    if inbound_invoices:
        doc.add_heading("Notas de Entrada", level=2)
        
        table_in = doc.add_table(rows=1, cols=5)
        table_in.style = 'Table Grid'
        
        # Cabeçalhos
        header_cells = table_in.rows[0].cells
        header_cells[0].text = "Número"
        header_cells[1].text = "Data Emissão"
        header_cells[2].text = "Fornecedor"
        header_cells[3].text = "Valor Total"
        header_cells[4].text = "Impostos"
        
        # Formatar cabeçalhos
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
        
        # Dados
        for invoice in sorted(inbound_invoices, key=lambda x: x.issue_date):
            row_cells = table_in.add_row().cells
            row_cells[0].text = f"{invoice.invoice_number}-{invoice.series}"
            row_cells[1].text = invoice.issue_date.strftime('%d/%m/%Y')
            row_cells[2].text = invoice.supplier.name if invoice.supplier else ""
            row_cells[3].text = f"R$ {float(invoice.total_value):,.2f}"
            row_cells[4].text = f"R$ {float(invoice.total_tax):,.2f}"
    
    # Resumo de impostos
    from services.tax_service import get_tax_summary_by_period
    tax_summary = get_tax_summary_by_period(start_date, end_date)
    
    if tax_summary:
        doc.add_page_break()
        doc.add_heading("Resumo de Impostos", level=1)
        
        table_tax = doc.add_table(rows=1, cols=2)
        table_tax.style = 'Table Grid'
        
        # Cabeçalhos
        header_cells = table_tax.rows[0].cells
        header_cells[0].text = "Imposto"
        header_cells[1].text = "Valor Total"
        
        # Formatar cabeçalhos
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
        
        # Dados
        for tax_type, data in tax_summary.items():
            row_cells = table_tax.add_row().cells
            row_cells[0].text = tax_type
            row_cells[1].text = f"R$ {data['total']:,.2f}"
            
            # Destacar valores negativos (créditos)
            if data['total'] < 0:
                row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    # Salvar o documento
    filename = os.path.join(output_dir, f"Relatorio_Financeiro_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    
    return filename
